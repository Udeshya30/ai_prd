import io
import json
import logging
from typing import AsyncGenerator

from fastapi import APIRouter, File, HTTPException, UploadFile
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, field_validator

from llama_runner import PRD_SECTIONS, generate_prd, stream_section, stream_thinking

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/generate-prd", tags=["prd"])


class PRDRequest(BaseModel):
    project_name: str
    problem: str
    features: str
    users: str
    goals: str

    @field_validator("project_name", "problem", "features", "users", "goals")
    @classmethod
    def not_empty(cls, v: str) -> str:
        if not v.strip():
            raise ValueError("Field must not be blank")
        if len(v) > 5000:
            raise ValueError("Field exceeds 5000 character limit")
        return v.strip()


# ---------------------------------------------------------------------------
# Non-streaming endpoint (kept for compatibility / upload flow)
# ---------------------------------------------------------------------------
@router.post("")
async def generate_prd_endpoint(request: PRDRequest):
    prompt = (
        "You are a senior product manager. Write a comprehensive PRD in markdown "
        "covering: objective, problem, target users, goals/KPIs, features, "
        "tech stack, architecture, dev phases & timeline, milestones, non-functional "
        "requirements, risks, and out-of-scope items.\n\n"
        f"Project: {request.project_name}\nProblem: {request.problem}\n"
        f"Features: {request.features}\nUsers: {request.users}\nGoals: {request.goals}\n\nPRD:\n"
    )
    logger.info("Generating PRD (single call) for: %s", request.project_name)
    try:
        result = generate_prd(prompt)
    except Exception as exc:
        logger.exception("LLM inference failed")
        raise HTTPException(status_code=500, detail="PRD generation failed.") from exc
    return {"prd": result}


# ---------------------------------------------------------------------------
# Orchestrated streaming: thinking → 12 fixed sections
#
# SSE event shapes:
#   {"type": "thinking_start"}
#   {"type": "thinking_token", "content": "..."}
#   {"type": "thinking_done", "analysis": "<full thinking text>"}
#   {"type": "section_start", "section": "Tech Stack", "index": 5, "total": 12}
#   {"type": "token", "content": "..."}
#   {"type": "section_done", "section": "Tech Stack"}
#   {"error": "message"}
#   [DONE]
# ---------------------------------------------------------------------------
@router.post("/stream")
async def stream_prd_orchestrated(request: PRDRequest):
    async def event_generator() -> AsyncGenerator[str, None]:
        def emit(obj: dict) -> str:
            return f"data: {json.dumps(obj)}\n\n"

        # ── Phase 1: Thinking / reasoning ──────────────────────────────────
        yield emit({"type": "thinking_start"})
        thinking_tokens: list[str] = []
        try:
            for token in stream_thinking(
                request.project_name, request.problem,
                request.features, request.users, request.goals,
            ):
                thinking_tokens.append(token)
                yield emit({"type": "thinking_token", "content": token})
        except Exception:
            logger.exception("Thinking phase failed")
            yield emit({"error": "Thinking phase failed. Continuing without analysis."})

        thinking = "".join(thinking_tokens)
        yield emit({"type": "thinking_done", "analysis": thinking})

        # ── Phase 2: One LLM call per section ──────────────────────────────
        total = len(PRD_SECTIONS)
        for idx, (key, title, max_tok) in enumerate(PRD_SECTIONS):
            yield emit({"type": "section_start", "section": title, "index": idx, "total": total})
            # Emit the heading so it appears immediately in the PRD
            yield emit({"type": "token", "content": f"## {title}\n\n"})
            try:
                for token in stream_section(
                    key, title,
                    request.project_name, request.problem,
                    request.features, request.users, request.goals,
                    thinking, max_tok,
                ):
                    yield emit({"type": "token", "content": token})
            except Exception:
                logger.exception("Section generation failed: %s", title)
                yield emit({"type": "token", "content": "\n\n*Error generating this section.*\n\n"})

            yield emit({"type": "section_done", "section": title})
            yield emit({"type": "token", "content": "\n\n---\n\n"})

        yield "data: [DONE]\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream")


# ---------------------------------------------------------------------------
# File upload endpoint
# ---------------------------------------------------------------------------
upload_router = APIRouter(prefix="/upload-requirements", tags=["upload"])


@upload_router.post("")
async def upload_file(file: UploadFile = File(...)):
    if file.filename is None or not (
        file.filename.endswith(".txt") or file.filename.endswith(".docx")
    ):
        raise HTTPException(status_code=422, detail="Only .txt or .docx files are supported.")

    content = await file.read()
    if len(content) > 5 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="File too large. Maximum size is 5 MB.")

    if file.filename.endswith(".txt"):
        text = content.decode("utf-8", errors="replace")
    else:
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            raise HTTPException(status_code=500, detail="python-docx not installed on server.")

    if not text.strip():
        raise HTTPException(status_code=422, detail="Uploaded file contains no readable text.")

    prompt = (
        "You are a senior product manager. Based on the following requirements, write a detailed PRD "
        "in markdown covering: executive summary, problem, target users, goals/KPIs, features, tech stack, "
        "architecture, development phases & timeline, milestones, non-functional requirements, risks, "
        "and out-of-scope items.\n\n"
        f"Requirements:\n{text}\n\nPRD:\n"
    )
    logger.info("Generating PRD from uploaded file: %s", file.filename)
    try:
        result = generate_prd(prompt)
    except Exception as exc:
        logger.exception("LLM inference failed for uploaded file")
        raise HTTPException(status_code=500, detail="PRD generation failed.") from exc
    return {"prd": result}



class PRDRequest(BaseModel):
    project_name: str
    problem: str
    features: str
    users: str
    goals: str

    @field_validator("project_name", "problem", "features", "users", "goals")
    @classmethod
    def not_empty(cls, v: str) -> str:
        if not v.strip():
            raise ValueError("Field must not be blank")
        if len(v) > 5000:
            raise ValueError("Field exceeds 5000 character limit")
        return v.strip()


def _build_prompt(
    project_name: str,
    problem: str,
    features: str,
    users: str,
    goals: str,
) -> str:
    return (
        "You are a senior product manager. Write a comprehensive Product Requirements Document (PRD) "
        "based on the details below. Use markdown with clear headings, bullet points, and tables where "
        "appropriate. The document must be at least 2000 words.\n\n"
        f"Project Name: {project_name}\n"
        f"Problem Statement: {problem}\n"
        f"Key Features: {features}\n"
        f"Target Users: {users}\n"
        f"Goals & Success Metrics: {goals}\n\n"
        "PRD:\n"
    )


@router.post("")
async def generate_prd_endpoint(request: PRDRequest):
    prompt = _build_prompt(
        request.project_name,
        request.problem,
        request.features,
        request.users,
        request.goals,
    )
    logger.info("Generating PRD for project: %s", request.project_name)
    try:
        result = generate_prd(prompt)
    except Exception as exc:
        logger.exception("LLM inference failed")
        raise HTTPException(status_code=500, detail="PRD generation failed. Check server logs.") from exc
    return {"prd": result}


@router.post("/stream")
async def stream_prd_endpoint(request: PRDRequest):
    """Server-Sent Events endpoint — streams tokens as they are generated."""
    prompt = _build_prompt(
        request.project_name,
        request.problem,
        request.features,
        request.users,
        request.goals,
    )
    logger.info("Streaming PRD for project: %s", request.project_name)

    async def event_generator() -> AsyncGenerator[str, None]:
        try:
            for token in stream_prd(prompt):
                yield f"data: {json.dumps({'token': token})}\n\n"
        except Exception:
            logger.exception("LLM streaming failed")
            yield f"data: {json.dumps({'error': 'Generation failed'})}\n\n"
        finally:
            yield "data: [DONE]\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream")


upload_router = APIRouter(prefix="/upload-requirements", tags=["upload"])


@upload_router.post("")
async def upload_file(file: UploadFile = File(...)):
    if file.filename is None or not (
        file.filename.endswith(".txt") or file.filename.endswith(".docx")
    ):
        raise HTTPException(status_code=422, detail="Only .txt or .docx files are supported.")

    content = await file.read()
    if len(content) > 5 * 1024 * 1024:  # 5 MB guard
        raise HTTPException(status_code=413, detail="File too large. Maximum size is 5 MB.")

    if file.filename.endswith(".txt"):
        text = content.decode("utf-8", errors="replace")
    else:
        try:
            from docx import Document  # lazy import — won't crash if not installed at startup
            doc = Document(io.BytesIO(content))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            raise HTTPException(status_code=500, detail="python-docx not installed on server.")

    if not text.strip():
        raise HTTPException(status_code=422, detail="Uploaded file contains no readable text.")

    prompt = (
        "You are a senior product manager. Based on the following requirements, write a detailed PRD "
        "in markdown format with at least 2000 words. Use clear headings, bullet points, and tables "
        "where appropriate.\n\n"
        f"Requirements:\n{text}\n\n"
        "PRD:\n"
    )
    logger.info("Generating PRD from uploaded file: %s", file.filename)
    try:
        result = generate_prd(prompt)
    except Exception as exc:
        logger.exception("LLM inference failed for uploaded file")
        raise HTTPException(status_code=500, detail="PRD generation failed. Check server logs.") from exc
    return {"prd": result}
